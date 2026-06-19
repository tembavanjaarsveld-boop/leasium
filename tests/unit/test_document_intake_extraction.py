import json
from io import BytesIO
from typing import Any

import httpx
import pytest
from docx import Document
from stewart.ai.document_intake import DocumentExtractionError, extract_document_file
from stewart.ai.lease_intake import _extract_document_text
from stewart.core.settings import Settings

INVOICE_TEXT = """PAYMENT ADVICE
To: SJI No 1 Pty Ltd
ABN: 97 656 552 699
Email: admin@skjproperty.com.au
Phone: +61 459 380 421
PO Box 5433
Brendale QLD 4500
Customer Gorilla Grind
Invoice Number INV-0331
Amount Due 0.00
Due Date 3 Jun 2026
TAX INVOICE
Gorilla Grind
Invoice Date
31 May 2026
Invoice Number
INV-0331
Reference
Gorilla Grind May 2026
Rent
ABN
97 656 552 699
SJI No 1 Pty Ltd
ABN: 97 656 552 699
Email:
admin@skjproperty.com.au
Description Quantity Unit Price GST Amount AUD
Rent - U1, B3 205 Leitchs Rd Brendale 1.00 3,833.33 10% 3,833.33
Rent - U3, B3 205 Leitchs Rd Brendale 1.00 4,083.33 10% 4,083.33
Subtotal 7,916.66
TOTAL GST 10% 791.66
TOTAL AUD 8,708.32
Less Amount Paid 8,708.32
AMOUNT DUE AUD 0.00
Due Date: 3 Jun 2026
Account name: SJI No 1 Pty Ltd
"""

INVOICE_LAYOUT_TEXT = "\n".join(
    [
        "Invoice Date                SJI No 1 Pty Ltd",
        "TAX INVOICE                 31 May 2026                 ABN: 97 656 552 699",
        "Invoice Number              Email:",
        "Gorilla Grind               INV-0331                    admin@skjproperty.com.au",
        "Reference",
        "Gorilla Grind May 2026",
        "Rent",
        "Description Quantity Unit Price GST Amount AUD",
        "Rent - U1, B3 205 Leitchs Rd Brendale 1.00 3,833.33 10% 3,833.33",
        "Rent - U3, B3 205 Leitchs Rd Brendale 1.00 4,083.33 10% 4,083.33",
        "TOTAL AUD                8,708.32",
        "Less Amount Paid               8,708.32",
        "AMOUNT DUE AUD                        0.00",
        "Due Date: 3 Jun 2026",
        "PAYMENT ADVICE",
        "Customer              Gorilla Grind",
        "To:        SJI No 1 Pty Ltd                                               Amount Enclosed",
        "ABN: 97 656 552 699",
        "Email: admin@skjproperty.com.au",
    ]
)


def _empty_invoice_extraction() -> dict[str, Any]:
    return {
        "document_type": "invoice_admin",
        "summary": (
            "Tax invoice for Gorilla Grind issued by SJI No 1 Pty Ltd for May "
            "2026 rent across two premises."
        ),
        "confidence": 0.88,
        "parties": [],
        "properties": [],
        "key_dates": [],
        "money_amounts": [],
        "tenancy_schedule": [],
        "obligations": [],
        "suggested_links": {
            "property_name": None,
            "tenant_name": None,
            "lease_reference": None,
        },
        "warnings": [],
        "missing_information": [],
        "proposed_actions": [],
        "inspection_findings": [],
    }


class _FakeOpenAIResponse:
    def __init__(
        self,
        *,
        body: dict[str, Any] | None = None,
        output_text: str | None = None,
        status_code: int | None = None,
    ) -> None:
        self._body = body
        self._output_text = (
            json.dumps(_empty_invoice_extraction()) if output_text is None else output_text
        )
        self._status_code = status_code

    def raise_for_status(self) -> None:
        if self._status_code is not None:
            request = httpx.Request("POST", "https://api.openai.com/v1/responses")
            response = httpx.Response(self._status_code, request=request)
            response.raise_for_status()
        return None

    def json(self) -> dict[str, Any]:
        if self._body is not None:
            return self._body
        return {
            "id": "resp_empty_invoice",
            "output": [
                {
                    "content": [
                        {
                            "type": "output_text",
                            "text": self._output_text,
                        }
                    ]
                }
            ],
        }


class _FakeHTTPClient:
    def __init__(
        self,
        *,
        timeout: float,
        response: _FakeOpenAIResponse | None = None,
        post_exception: Exception | None = None,
    ) -> None:
        self.timeout = timeout
        self._response = response or _FakeOpenAIResponse()
        self._post_exception = post_exception

    def __enter__(self) -> "_FakeHTTPClient":
        return self

    def __exit__(self, *args: object) -> None:
        return None

    def post(self, *args: object, **kwargs: object) -> _FakeOpenAIResponse:
        if self._post_exception is not None:
            raise self._post_exception
        return self._response


def _patch_http_client(
    monkeypatch: Any,
    *,
    response: _FakeOpenAIResponse | None = None,
    post_exception: Exception | None = None,
) -> None:
    def fake_client(*, timeout: float) -> _FakeHTTPClient:
        return _FakeHTTPClient(
            timeout=timeout,
            response=response,
            post_exception=post_exception,
        )

    monkeypatch.setattr("stewart.ai.document_intake.httpx.Client", fake_client)


def test_extract_document_file_requires_openai_api_key() -> None:
    with pytest.raises(DocumentExtractionError, match="OpenAI API key is not configured"):
        extract_document_file(
            file_data=b"notice",
            filename="notice.txt",
            content_type="text/plain",
            settings=Settings(openai_api_key=""),
        )


@pytest.mark.parametrize("status_code", [429, 503])
def test_extract_document_file_reports_openai_status_errors(
    monkeypatch: Any,
    status_code: int,
) -> None:
    _patch_http_client(
        monkeypatch,
        response=_FakeOpenAIResponse(status_code=status_code),
    )

    with pytest.raises(
        DocumentExtractionError,
        match=f"OpenAI extraction request failed with status {status_code}",
    ):
        extract_document_file(
            file_data=b"notice",
            filename="notice.txt",
            content_type="text/plain",
            settings=Settings(openai_api_key="sk-test"),
        )


def test_extract_document_file_reports_openai_timeouts(monkeypatch: Any) -> None:
    request = httpx.Request("POST", "https://api.openai.com/v1/responses")
    _patch_http_client(
        monkeypatch,
        post_exception=httpx.TimeoutException("timed out", request=request),
    )

    with pytest.raises(
        DocumentExtractionError,
        match="OpenAI extraction request timed out",
    ):
        extract_document_file(
            file_data=b"notice",
            filename="notice.txt",
            content_type="text/plain",
            settings=Settings(openai_api_key="sk-test"),
        )


def test_extract_document_file_reports_malformed_openai_json(monkeypatch: Any) -> None:
    _patch_http_client(
        monkeypatch,
        response=_FakeOpenAIResponse(output_text="{not-json"),
    )

    with pytest.raises(
        DocumentExtractionError,
        match="OpenAI response was not valid JSON",
    ):
        extract_document_file(
            file_data=b"notice",
            filename="notice.txt",
            content_type="text/plain",
            settings=Settings(openai_api_key="sk-test"),
        )


def test_extract_document_file_rejects_missing_top_level_fields(
    monkeypatch: Any,
) -> None:
    _patch_http_client(
        monkeypatch,
        response=_FakeOpenAIResponse(
            output_text=json.dumps(
                {
                    "document_type": "notice",
                    "summary": "A notice for operator review.",
                    "confidence": 0.8,
                    "suggested_links": {
                        "property_name": None,
                        "tenant_name": None,
                        "lease_reference": None,
                    },
                }
            ),
        ),
    )

    with pytest.raises(
        DocumentExtractionError,
        match="OpenAI extraction was missing required fields",
    ):
        extract_document_file(
            file_data=b"notice",
            filename="notice.txt",
            content_type="text/plain",
            settings=Settings(openai_api_key="sk-test"),
        )


def test_extract_document_file_preserves_inspection_report_type(monkeypatch: Any) -> None:
    extraction = _empty_invoice_extraction()
    extraction["document_type"] = "inspection_report"
    extraction["summary"] = "Inspection report for reviewed maintenance findings."
    _patch_http_client(
        monkeypatch,
        response=_FakeOpenAIResponse(output_text=json.dumps(extraction)),
    )

    extracted, _response_id = extract_document_file(
        file_data=b"inspection",
        filename="inspection.txt",
        content_type="text/plain",
        settings=Settings(openai_api_key="sk-test"),
    )

    assert extracted["document_type"] == "inspection_report"


def test_extract_document_file_supplements_empty_invoice_from_source_text(
    monkeypatch: Any,
) -> None:
    monkeypatch.setattr("stewart.ai.document_intake.httpx.Client", _FakeHTTPClient)

    extracted, response_id = extract_document_file(
        file_data=INVOICE_TEXT.encode(),
        filename="Invoice INV-0331.txt",
        content_type="text/plain",
        settings=Settings(openai_api_key="sk-test"),
    )

    assert response_id == "resp_empty_invoice"
    assert extracted["document_type"] == "invoice_admin"
    assert extracted["suggested_links"]["tenant_name"] == "Gorilla Grind"
    assert extracted["suggested_links"]["property_name"] == "205 Leitchs Rd Brendale"
    assert [party["name"] for party in extracted["parties"]] == [
        "Gorilla Grind",
        "SJI No 1 Pty Ltd",
    ]
    assert extracted["properties"][0]["address"] == "205 Leitchs Rd Brendale"
    assert extracted["properties"][0]["unit_label"] == "U1, B3; U3, B3"
    assert extracted["key_dates"] == [
        {
            "label": "Invoice date",
            "date": "2026-05-31",
            "confidence": 0.9,
            "source_hint": "Invoice Date 31 May 2026",
        },
        {
            "label": "Payment due",
            "date": "2026-06-03",
            "confidence": 0.9,
            "source_hint": "Due Date 3 Jun 2026",
        },
    ]
    assert extracted["money_amounts"][0] == {
        "label": "Total rent invoice including GST",
        "amount": 8708.32,
        "currency": "AUD",
        "frequency": "monthly",
        "confidence": 0.9,
        "source_hint": "TOTAL AUD 8,708.32",
    }
    assert extracted["proposed_actions"][0]["action"] == "prepare_billing_review"
    assert any("Amount Due AUD 0.00" in warning for warning in extracted["warnings"])


def test_extract_document_file_supplements_column_layout_invoice_text(
    monkeypatch: Any,
) -> None:
    monkeypatch.setattr("stewart.ai.document_intake.httpx.Client", _FakeHTTPClient)

    extracted, _response_id = extract_document_file(
        file_data=INVOICE_LAYOUT_TEXT.encode(),
        filename="Invoice INV-0331.txt",
        content_type="text/plain",
        settings=Settings(openai_api_key="sk-test"),
    )

    assert extracted["parties"][1]["name"] == "SJI No 1 Pty Ltd"
    assert extracted["properties"][0]["invoice_reference"] == "INV-0331"
    assert extracted["properties"][0]["invoice_issuer_name"] == "SJI No 1 Pty Ltd"
    assert extracted["money_amounts"][0]["frequency"] == "monthly"


def test_extract_document_text_reads_docx_tables() -> None:
    # Commercial leases (QTR Form 7/20) keep parties and dates in tables. The
    # reader must include table cells, not just paragraphs, or the model never
    # sees the tenant/landlord names or the lease term.
    doc = Document()
    doc.add_paragraph("QUEENSLAND TITLES REGISTRY — LEASE")
    table = doc.add_table(rows=0, cols=2)
    for label, value in [
        ("Lessor (Landlord)", "SKJ Property"),
        ("Lessee (Tenant)", "SKJ Capital"),
        ("Commencement date", "01/08/2022"),
        ("Expiry date", "31/07/2027"),
    ]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    buffer = BytesIO()
    doc.save(buffer)

    text = _extract_document_text(buffer.getvalue(), "lease.docx", None)

    assert text is not None
    assert "QUEENSLAND TITLES REGISTRY" in text
    assert "Lessee (Tenant) | SKJ Capital" in text
    assert "SKJ Property" in text
    assert "01/08/2022" in text
    assert "31/07/2027" in text
